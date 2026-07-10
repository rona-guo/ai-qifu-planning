# -*- coding: utf-8 -*-
"""
生成《AI时代企业数字化转型实战培训体系》通用版
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return heading

def create_outline_doc():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI时代企业数字化转型实战培训体系")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 通用版")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("「商业为本 · 场景贯穿 · 成果交付 · 价值验证」\n\n一套适用于各行业的数字化转型实战培训体系")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ===== 核心理念 =====
    add_heading(doc, "核心理念", level=1, color=(0, 51, 102))
    
    doc.add_paragraph("本培训体系以「商业为本、场景贯穿、成果交付、价值验证」为核心理念，采用行动学习法，让学员代入自己的企业场景，在解决实际问题的过程中掌握数字化转型的完整方法论。")

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run("\n" + "="*50 + "\n")
    run.font.size = Pt(12)
    run = formula.add_run("商业成功 = 理解商业本质 × 选对商业模式 × 用好工具（数字化+AI）\n")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    run = formula.add_run("="*50 + "\n")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ===== 培训体系全景图 =====
    add_heading(doc, "培训体系全景图：四环驱动模型", level=1, color=(0, 51, 102))

    doc.add_paragraph("本培训体系采用「四环驱动」模型，贯穿从诊断到验证的完整闭环：")

    modules = [
        ("模块一", "商业诊断", "现状诊断与机会识别", "发现核心问题"),
        ("模块二", "方案设计", "转型方案与路线图", "制定改善策略"),
        ("模块三", "系统建设", "数字化系统与方法", "实现数据驱动"),
        ("模块四", "变革落地", "组织变革与运营", "推动持续改进"),
        ("彩蛋", "价值验证", "成果交付与验证", "量化转型价值"),
    ]

    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'

    headers = ["模块", "主题", "核心动作", "产出"]
    for i, h in enumerate(headers):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (m, t, a, p) in enumerate(modules, 1):
        table2.rows[row_idx].cells[0].text = m
        table2.rows[row_idx].cells[1].text = t
        table2.rows[row_idx].cells[2].text = a
        table2.rows[row_idx].cells[3].text = p

    doc.add_paragraph()

    # ===== 学员自带场景说明 =====
    doc.add_paragraph("【学员自带场景】学员在培训前提交自己企业的1-2个核心业务痛点，培训过程中导师将指导学员结合所学方法论分析自己的企业场景，形成可落地的改善方案。")

    doc.add_page_break()

    # ===== 模块一详细内容 =====
    add_heading(doc, "模块一：商业诊断", level=1, color=(0, 102, 153))

    add_heading(doc, "1.1 培训目标", level=2)
    doc.add_paragraph("掌握企业现状诊断的核心方法论")
    doc.add_paragraph("能够独立分析业务痛点的根因")
    doc.add_paragraph("识别数字化转型的核心机会点")

    add_heading(doc, "1.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】企业诊断三板斧", style='Intense Quote')
    doc.add_paragraph("")
    doc.add_paragraph("第一板斧：商业模式画布分析")
    doc.add_paragraph("  - 梳理企业价值主张、客户群体、渠道、核心资源")
    doc.add_paragraph("  - 识别商业模式中的数字化机会点")
    doc.add_paragraph("")
    doc.add_paragraph("第二板斧：价值链诊断")
    doc.add_paragraph("  - 绘制企业核心价值链")
    doc.add_paragraph("  - 识别各环节的效率损耗点")
    doc.add_paragraph("  - 定位数字化可以发力的关键环节")
    doc.add_paragraph("")
    doc.add_paragraph("第三板斧：根因分析五问法")
    doc.add_paragraph("  - 这个问题背后的问题是什么？")
    doc.add_paragraph("  - 数据支撑还是经验判断？")
    doc.add_paragraph("  - 是流程问题还是系统问题？")
    doc.add_paragraph("  - 是组织问题还是文化问题？")
    doc.add_paragraph("  - 不解决会有什么后果？")

    add_heading(doc, "1.3 诊断工具包", level=2)
    
    tools_table = doc.add_table(rows=5, cols=2)
    tools_table.style = 'Table Grid'
    tool_headers = ["工具", "用途"]
    for i, h in enumerate(tool_headers):
        cell = tools_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tools = [
        ("商业模式画布", "梳理企业商业模式，识别数字化机会"),
        ("价值链分析图", "绘制核心价值链，定位效率损耗点"),
        ("根因分析模板", "5Why分析法，找到问题的根本原因"),
        ("痛点优先级矩阵", "评估痛点影响度和解决紧迫性"),
    ]
    for row_idx, data in enumerate(tools, 1):
        for col_idx, val in enumerate(data):
            tools_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "1.4 小组练习", level=2)
    doc.add_paragraph("学员运用所学方法论分析自己企业的核心痛点，形成诊断报告")

    add_heading(doc, "1.5 核心产出", level=2)
    doc.add_paragraph("[OK] 企业现状诊断报告（学员自带场景）")
    doc.add_paragraph("[OK] 痛点根因分析报告")
    doc.add_paragraph("[OK] 数字化转型机会地图（TOP10）")

    doc.add_page_break()

    # ===== 模块二详细内容 =====
    add_heading(doc, "模块二：方案设计", level=1, color=(0, 102, 153))

    add_heading(doc, "2.1 培训目标", level=2)
    doc.add_paragraph("能够设计完整的数字化转型方案")
    doc.add_paragraph("掌握转型路线图的制定方法")
    doc.add_paragraph("学会用「价值承诺」锁定改善目标")

    add_heading(doc, "2.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】转型方案设计三步法", style='Intense Quote')
    doc.add_paragraph("")
    doc.add_paragraph("第一步：锁定核心问题")
    doc.add_paragraph("  - 从TOP10机会中筛选TOP3核心机会")
    doc.add_paragraph("  - 评估各机会的ROI和实施复杂度")
    doc.add_paragraph("  - 确定优先级，聚焦突破")
    doc.add_paragraph("")
    doc.add_paragraph("第二步：设计解决方案")
    doc.add_paragraph("  - 流程优化：先优化流程，再上系统")
    doc.add_paragraph("  - 数字工具：选择合适的数字化工具组合")
    doc.add_paragraph("  - 组织配套：配套的考核激励和人员能力")
    doc.add_paragraph("")
    doc.add_paragraph("第三步：承诺价值目标")
    doc.add_paragraph("  - 量化每个改善项的预期收益")
    doc.add_paragraph("  - 明确验收标准和验证方法")
    doc.add_paragraph("  - 制定里程碑计划")

    doc.add_paragraph("【案例库】各行业转型方案参考", style='Intense Quote')
    doc.add_paragraph("制造业案例：订单交付效率提升方案")
    doc.add_paragraph("零售业案例：全渠道会员运营方案")
    doc.add_paragraph("服务业案例：服务流程标准化方案")
    doc.add_paragraph("（导师将根据学员行业补充对应案例）")

    add_heading(doc, "2.3 转型路线图设计", level=2)
    
    roadmap_table = doc.add_table(rows=5, cols=4)
    roadmap_table.style = 'Table Grid'
    rm_headers = ["阶段", "时间", "核心任务", "关键里程碑"]
    for i, h in enumerate(rm_headers):
        cell = roadmap_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    roadmap = [
        ("POC验证期", "1-2月", "选择1个场景试点，验证方案可行性", "试点场景效率+20%"),
        ("快速推广期", "3-6月", "推广至核心业务场景", "核心场景全面覆盖"),
        ("深度优化期", "7-12月", "数据驱动决策，AI智能预测", "形成数据资产"),
        ("持续运营期", "次年+", "内部能力固化，持续迭代", "ROI持续验证"),
    ]
    for row_idx, data in enumerate(roadmap, 1):
        for col_idx, val in enumerate(data):
            roadmap_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "2.4 小组练习", level=2)
    doc.add_paragraph("学员针对自己企业的TOP1痛点，设计完整的转型方案和路线图")

    add_heading(doc, "2.5 核心产出", level=2)
    doc.add_paragraph("[OK] 企业数字化转型方案v1.0（学员自带场景）")
    doc.add_paragraph("[OK] 12个月数字化转型路线图")
    doc.add_paragraph("[OK] 转型价值承诺书")

    doc.add_page_break()

    # ===== 模块三详细内容 =====
    add_heading(doc, "模块三：系统建设", level=1, color=(0, 102, 153))

    add_heading(doc, "3.1 培训目标", level=2)
    doc.add_paragraph("理解企业数字化系统的核心架构")
    doc.add_paragraph("掌握系统选型的核心方法")
    doc.add_paragraph("能够设计数据看板和预警规则")

    add_heading(doc, "3.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】企业数字化系统架构", style='Intense Quote')
    doc.add_paragraph("")
    doc.add_paragraph("           ┌─────────────────────────────────────┐")
    doc.add_paragraph("           │           业务应用层                │")
    doc.add_paragraph("           │   CRM / ERP / SCM / MES / OA...    │")
    doc.add_paragraph("           └─────────────┬───────────────────────┘")
    doc.add_paragraph("                         │")
    doc.add_paragraph("           ┌─────────────┴───────────────────────┐")
    doc.add_paragraph("           │           数据中台层                 │")
    doc.add_paragraph("           │    主数据 / 数据湖 / API网关        │")
    doc.add_paragraph("           └─────────────┬───────────────────────┘")
    doc.add_paragraph("                         │")
    doc.add_paragraph("           ┌─────────────┴───────────────────────┐")
    doc.add_paragraph("           │           基础设施层                │")
    doc.add_paragraph("           │    云服务器 / 网络安全 / 存储       │")
    doc.add_paragraph("           └─────────────────────────────────────┘")

    doc.add_paragraph("【核心原则】系统选型三问法", style='Intense Quote')
    doc.add_paragraph("")
    doc.add_paragraph("第一问：这个系统解决什么业务问题？")
    doc.add_paragraph("  - 不要为了上系统而上系统")
    doc.add_paragraph("  - 先明确要解决的问题，再选系统")
    doc.add_paragraph("")
    doc.add_paragraph("第二问：这个系统能被员工真正使用吗？")
    doc.add_paragraph("  - 易用性 > 功能性")
    doc.add_paragraph("  - 再好的系统，没人用等于零")
    doc.add_paragraph("")
    doc.add_paragraph("第三问：这个系统能和企业现有系统集成吗？")
    doc.add_paragraph("  - 数据孤岛是数字化最大敌人")
    doc.add_paragraph("  - 选型时必须评估集成能力")

    doc.add_paragraph("【工具推荐】主流数字化工具一览", style='Intense Quote')
    
    tool_cat_table = doc.add_table(rows=9, cols=3)
    tool_cat_table.style = 'Table Grid'
    tc_headers = ["类别", "代表产品", "适用场景"]
    for i, h in enumerate(tc_headers):
        cell = tool_cat_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tool_cats = [
        ("协同办公", "钉钉/飞书/企业微信", "日常沟通、审批、考勤"),
        ("ERP系统", "用友/金蝶/SAP", "财务、采购、库存管理"),
        ("CRM系统", "销售易/纷享销客", "客户管理、销售跟进"),
        ("MES系统", "西门子/鼎捷/慧程", "生产执行、工序跟踪"),
        ("数据分析", "帆软/Power BI/Tableau", "数据可视化、自助分析"),
        ("AI工具", "ChatGPT/文心/通义", "内容生成、智能客服、数据分析"),
        ("低代码", "简道云/氚云/宜搭", "轻量级流程应用搭建"),
        ("项目管理", "Teambition/禅道/Jira", "任务管理、进度跟踪"),
    ]
    for row_idx, data in enumerate(tool_cats, 1):
        for col_idx, val in enumerate(data):
            tool_cat_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "3.3 AI工具实操", level=2)
    doc.add_paragraph("使用AI工具进行会议纪要自动生成")
    doc.add_paragraph("使用AI工具进行数据趋势分析和预测")
    doc.add_paragraph("使用AI工具构建智能问答知识库")
    doc.add_paragraph("使用AI工具辅助报告撰写和方案设计")

    add_heading(doc, "3.4 核心产出", level=2)
    doc.add_paragraph("[OK] 企业数字化系统需求文档")
    doc.add_paragraph("[OK] 系统选型评估表（含评估维度）")
    doc.add_paragraph("[OK] 管理驾驶舱原型（Excel版）")

    doc.add_page_break()

    # ===== 模块四详细内容 =====
    add_heading(doc, "模块四：变革落地", level=1, color=(0, 102, 153))

    add_heading(doc, "4.1 培训目标", level=2)
    doc.add_paragraph("掌握组织变革推动的关键方法")
    doc.add_paragraph("能够建立内部数字化运营团队")
    doc.add_paragraph("设计有效的激励机制和考核体系")

    add_heading(doc, "4.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】数字化转型变革管理四步法", style='Intense Quote')
    doc.add_paragraph("")
    doc.add_paragraph("第一步：领导力驱动")
    doc.add_paragraph("  - 高层挂帅，成立数字化委员会")
    doc.add_paragraph("  - 明确一把手工程，由CEO或COO亲自推动")
    doc.add_paragraph("")
    doc.add_paragraph("第二步：能力建设")
    doc.add_paragraph("  - 培养内部数字化专员团队")
    doc.add_paragraph("  - 建立数字化技能认证体系")
    doc.add_paragraph("  - 引入外部专家辅导机制")
    doc.add_paragraph("")
    doc.add_paragraph("第三步：流程嵌入")
    doc.add_paragraph("  - 将数字化工具使用嵌入业务流程")
    doc.add_paragraph("  - 建立系统使用SLA和考核机制")
    doc.add_paragraph("  - 定期检视系统使用率和效果")
    doc.add_paragraph("")
    doc.add_paragraph("第四步：文化固化")
    doc.add_paragraph("  - 将数据驱动写入公司价值观")
    doc.add_paragraph("  - 树立数字化转型标杆人物")
    doc.add_paragraph("  - 建立经验分享和复盘机制")

    add_heading(doc, "4.3 常见变革阻力及应对策略", level=2)
    
    resist_table = doc.add_table(rows=5, cols=2)
    resist_table.style = 'Table Grid'
    r_headers = ["阻力类型", "应对策略"]
    for i, h in enumerate(r_headers):
        cell = resist_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    resistances = [
        ("人员抵触", "充分沟通、渐进式推进、快速胜利"),
        ("部门墙", "设立跨部门项目组、共享KPI"),
        ("资源不足", "聚焦核心场景、争取高层支持"),
        ("能力缺乏", "系统培训、引入外部顾问"),
    ]
    for row_idx, data in enumerate(resistances, 1):
        for col_idx, val in enumerate(data):
            resist_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "4.4 内部讲师认证", level=2)
    doc.add_paragraph("选拔学员成为企业内部数字化讲师")
    doc.add_paragraph("颁发「数字化内部认证讲师」证书")
    doc.add_paragraph("建立企业内部数字化知识库")

    add_heading(doc, "4.5 核心产出", level=2)
    doc.add_paragraph("[OK] 企业数字化变革落地方案")
    doc.add_paragraph("[OK] 内部数字化运营手册")
    doc.add_paragraph("[OK] 数字化考核激励方案")

    doc.add_page_break()

    # ===== 模块五详细内容（彩蛋） =====
    add_heading(doc, "模块五（彩蛋）：价值验证", level=1, color=(0, 102, 153))

    add_heading(doc, "5.1 培训目标", level=2)
    doc.add_paragraph("理解数字化转型价值验证的方法论")
    doc.add_paragraph("能够设计并追踪转型ROI")
    doc.add_paragraph("掌握向管理层汇报转型成果的技巧")

    add_heading(doc, "5.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】转型价值验证模型", style='Intense Quote')
    doc.add_paragraph("")
    doc.add_paragraph("转型ROI = (效率提升收益 + 成本降低收益 + 收入增长收益 - 转型投入) / 转型投入 x 100%")

    doc.add_paragraph("【案例】不同场景的价值测算参考", style='Intense Quote')
    
    value_table = doc.add_table(rows=6, cols=3)
    value_table.style = 'Table Grid'
    val_headers = ["改善场景", "典型收益", "测算方式"]
    for i, h in enumerate(val_headers):
        cell = value_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    values = [
        ("订单交付效率", "减少加班费、降低客户流失", "交付准时率提升百分点 x 基数"),
        ("客户运营效率", "提升复购率、增加转介绍", "客户满意度提升 x 客户生命周期价值"),
        ("生产效率提升", "减少在制品库存、降低报废", "良品率提升 x 产量 x 单件成本"),
        ("人力成本优化", "减少重复劳动、提升人效", "节省工时 x 人力成本"),
        ("决策质量提升", "减少拍脑袋决策的损失", "历史决策失误损失 x 改善比例"),
    ]
    for row_idx, data in enumerate(values, 1):
        for col_idx, val in enumerate(data):
            value_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "5.3 价值可视化呈现", level=2)
    doc.add_paragraph("管理驾驶舱：实时展示核心业务指标")
    doc.add_paragraph("价值追踪表：月度展示转型ROI趋势")
    doc.add_paragraph("管理层汇报：3页精华版，直击转型价值")

    add_heading(doc, "5.4 核心产出", level=2)
    doc.add_paragraph("[OK] 企业管理驾驶舱模板（Excel可操作版）")
    doc.add_paragraph("[OK] 月度转型价值验证报告模板")
    doc.add_paragraph("[OK] 管理层汇报PPT模板（3页精华版）")

    doc.add_page_break()

    # ===== 培训安排 =====
    add_heading(doc, "培训安排建议", level=1, color=(0, 51, 102))

    schedule_table = doc.add_table(rows=7, cols=4)
    schedule_table.style = 'Table Grid'
    sch_headers = ["模块", "建议时长", "培训形式", "产出数量"]
    for i, h in enumerate(sch_headers):
        cell = schedule_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    schedules = [
        ("模块一", "1天", "理论+实操+作业点评", "3项产出"),
        ("模块二", "1天", "工作坊+方案路演", "3项产出"),
        ("模块三", "1天", "演示+实操演练", "3项产出"),
        ("模块四", "0.5天", "工作坊+讨论", "3项产出"),
        ("彩蛋模块", "0.5天", "案例+模板", "3项产出"),
        ("综合演练", "半天（可选）", "沙盘模拟+答辩", "综合方案"),
    ]
    for row_idx, data in enumerate(schedules, 1):
        for col_idx, val in enumerate(data):
            schedule_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph("【标准版】建议总时长：4天（可分2周执行，每周2天）")
    doc.add_paragraph("【紧凑版】建议总时长：2天（聚焦核心模块）")
    doc.add_paragraph("学员人数：建议20-40人/期")
    doc.add_paragraph("课后辅导：提供30天线上答疑服务")

    doc.add_paragraph()
    doc.add_paragraph("【增值服务】（可选）")
    doc.add_paragraph("  - 陪跑服务：3-6个月驻场辅导")
    doc.add_paragraph("  - 系统选型：协助对接供应商")
    doc.add_paragraph("  - 行业定制：根据行业特点补充案例")

    doc.add_page_break()

    # ===== 附录：核心工具清单 =====
    add_heading(doc, "附录：培训核心工具清单", level=1, color=(0, 51, 102))

    tools_table = doc.add_table(rows=11, cols=3)
    tools_table.style = 'Table Grid'
    tool_headers = ["类别", "工具名称", "用途"]
    for i, h in enumerate(tool_headers):
        cell = tools_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    all_tools = [
        ("商业模式", "Canvanizer", "绘制商业模式画布"),
        ("流程设计", "Visio/Miro", "绘制业务流程图"),
        ("数据分析", "Excel/Power BI", "数据分析与可视化"),
        ("AI工具", "ChatGPT/文心一言", "内容生成、辅助分析"),
        ("协同办公", "飞书/钉钉/企业微信", "日常沟通、任务管理"),
        ("文档协作", "腾讯文档/飞书文档", "知识沉淀与共享"),
        ("项目管理", "Teambition", "任务跟踪与协作"),
        ("原型设计", "Axure/Figma", "系统原型设计"),
        ("演示汇报", "PPT模板包", "管理层汇报材料"),
        ("价值追踪", "Excel看板模板", "ROI追踪与验证"),
    ]
    for row_idx, data in enumerate(all_tools, 1):
        for col_idx, val in enumerate(data):
            tools_table.rows[row_idx].cells[col_idx].text = val

    # 保存文档
    output_path = r"d:\workboddy 培训设计\AI时代企业数字化转型实战培训体系_通用版.docx"
    doc.save(output_path)
    print(f"[OK] 文档已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    create_outline_doc()
