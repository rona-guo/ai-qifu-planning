# -*- coding: utf-8 -*-
"""
生成《AI时代企业数字化转型实战培训体系》XX制造企业场景版
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1, color=None):
    """添加标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return heading

def create_outline_doc():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ===== 封面页 =====
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI时代企业数字化转型实战培训体系")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— XX制造企业场景版")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()
    doc.add_paragraph()

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("以XX制造企业为贯穿案例\n从订单交付效率诊断到价值验证的完整培训体系")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ===== 核心理念 =====
    add_heading(doc, "核心理念", level=1, color=(0, 51, 102))
    
    doc.add_paragraph("本培训体系以「商业为本、场景贯穿、成果交付、价值验证」为核心理念，采用XX制造企业订单交付效率提升作为贯穿场景，让学员在真实企业案例中掌握数字化转型的完整方法论。")

    # 核心公式框
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run("\n" + "="*50 + "\n")
    run.font.size = Pt(12)
    run = formula.add_run("商业成功 = 理解商业本质 × 选对商业模式 × 用好工具（数字化+AI）\n")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    run = formula.add_run("="*50 + "\n")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ===== 案例企业背景 =====
    add_heading(doc, "贯穿案例：XX制造企业", level=1, color=(0, 51, 102))

    # 企业信息表格
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    info_data = [
        ("企业名称", "XX制造企业（化名：鑫智制造）"),
        ("企业规模", "年营收5亿元，员工800人，年产2000+SKU"),
        ("主营产品", "精密零部件制造，非标定制为主"),
        ("渠道结构", "直销60%、经销商30%、外贸10%"),
        ("当前痛点", "订单交付准时率72%、加班成本高、客户流失率15%"),
        ("培训目标", "交付准时率提升至95%、加班成本降低40%、客户流失率降至5%"),
    ]

    for i, (label, value) in enumerate(info_data):
        cell = table.rows[i].cells[0]
        cell.text = label
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "E8F4FD")
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ===== 培训体系全景图 =====
    add_heading(doc, "培训体系全景图：四环驱动模型", level=1, color=(0, 51, 102))

    doc.add_paragraph("本培训体系采用「四环驱动」模型，以订单交付效率提升为核心，贯穿从诊断到验证的完整闭环：")

    modules = [
        ("模块一", "商业诊断", "订单交付现状诊断", "发现效率瓶颈"),
        ("模块二", "方案设计", "转型方案与路线图", "制定改善策略"),
        ("模块三", "系统建设", "数字化系统搭建", "实现数据驱动"),
        ("模块四", "变革落地", "组织变革与运营", "推动持续改进"),
        ("彩蛋", "价值验证", "成果交付与验证", "量化转型价值"),
    ]

    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'

    headers = ["模块", "主题", "核心动作", "产出"]
    for i, h in enumerate(headers):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (m, t, a, p) in enumerate(modules, 1):
        table2.rows[row_idx].cells[0].text = m
        table2.rows[row_idx].cells[1].text = t
        table2.rows[row_idx].cells[2].text = a
        table2.rows[row_idx].cells[3].text = p

    doc.add_page_break()

    # ===== 模块一详细内容 =====
    add_heading(doc, "模块一：商业诊断", level=1, color=(0, 102, 153))

    add_heading(doc, "1.1 培训目标", level=2)
    doc.add_paragraph("掌握订单交付效率诊断的核心方法论")
    doc.add_paragraph("能够独立分析交付瓶颈的根因")
    doc.add_paragraph("识别鑫智制造的核心改善机会点")

    add_heading(doc, "1.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】订单交付诊断方法论", style='Intense Quote')
    doc.add_paragraph("交付效率 = 需求准确率 × 计划达成率 × 生产执行率 × 物流及时率")
    doc.add_paragraph("")
    doc.add_paragraph("诊断维度：")
    doc.add_paragraph("  - 需求端：订单需求准确率、历史订单交付分析")
    doc.add_paragraph("  - 计划端：排产合理性、产能利用率瓶颈")
    doc.add_paragraph("  - 生产端：工序节拍、良品率、在制品库存")
    doc.add_paragraph("  - 物流端：发货及时率、物流异常统计")

    doc.add_paragraph("【实战演练】鑫智制造订单交付现状分析", style='Intense Quote')

    # 交付现状表格
    analysis_table = doc.add_table(rows=6, cols=4)
    analysis_table.style = 'Table Grid'
    an_headers = ["环节", "当前指标", "行业标杆", "差距"]
    for i, h in enumerate(an_headers):
        cell = analysis_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    analysis_data = [
        ("订单需求准确率", "68%", "92%", "差距24%"),
        ("生产计划达成率", "75%", "95%", "差距20%"),
        ("一次良品率", "82%", "98%", "差距16%"),
        ("发货及时率", "72%", "98%", "差距26%"),
        ("整体交付准时率", "72%", "95%", "差距23%"),
    ]
    for row_idx, data in enumerate(analysis_data, 1):
        for col_idx, val in enumerate(data):
            analysis_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "1.3 根因分析：为什么交付准时率只有72%？", level=2)
    
    cause_table = doc.add_table(rows=5, cols=2)
    cause_table.style = 'Table Grid'
    cause_headers = ["根因类别", "具体表现"]
    for i, h in enumerate(cause_headers):
        cell = cause_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    causes = [
        ("需求端混乱", "订单变更频繁（35%订单有变更）、BOM不准确、技术交底不清"),
        ("计划排产粗放", "依赖经验排产、产能测算不准、紧急插单多（占比40%）"),
        ("车间执行低效", "工序等待时间长、换模时间长、质量返工多"),
        ("信息不透明", "订单进度靠电话追问、仓库配料不及时、数据滞后24h+"),
    ]
    for row_idx, data in enumerate(causes, 1):
        for col_idx, val in enumerate(data):
            cause_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "1.4 小组练习", level=2)
    doc.add_paragraph("分组绘制鑫智制造「订单交付全链路流程图」，标注每个环节的效率损耗点")

    add_heading(doc, "1.5 核心产出", level=2)
    doc.add_paragraph("[OK] 鑫智制造订单交付现状诊断报告")
    doc.add_paragraph("[OK] 交付效率瓶颈根因分析报告")
    doc.add_paragraph("[OK] 转型机会地图（TOP10改善机会）")

    doc.add_page_break()

    # ===== 模块二详细内容 =====
    add_heading(doc, "模块二：方案设计", level=1, color=(0, 102, 153))

    add_heading(doc, "2.1 培训目标", level=2)
    doc.add_paragraph("能够设计完整的订单交付改善方案")
    doc.add_paragraph("掌握数字化转型路线图的制定方法")
    doc.add_paragraph("学会用「价值承诺」锁定改善目标")

    add_heading(doc, "2.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】转型方案设计三步法", style='Intense Quote')
    doc.add_paragraph("第一步：锁定核心问题 - 聚焦TOP3改善机会")
    doc.add_paragraph("第二步：设计解决方案 - 流程优化+数字工具+组织配套")
    doc.add_paragraph("第三步：承诺价值目标 - 量化ROI和验收标准")

    doc.add_paragraph("【实战演练】鑫智制造转型方案设计", style='Intense Quote')

    solution_table = doc.add_table(rows=5, cols=4)
    solution_table.style = 'Table Grid'
    sol_headers = ["改善机会", "解决方案", "数字化工具", "预期提升"]
    for i, h in enumerate(sol_headers):
        cell = solution_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    solutions = [
        ("订单变更多", "建立ECN变更管控流程", "PLM变更管理模块", "变更频率-50%"),
        ("排产靠经验", "引入APS智能排产系统", "APS高级计划系统", "计划达成率+15%"),
        ("信息不透明", "部署MES+实时看板", "MES制造执行系统", "信息延迟<1h"),
        ("质量返工多", "引入SPC质量预警", "SPC统计过程控制", "良品率+10%"),
    ]
    for row_idx, data in enumerate(solutions, 1):
        for col_idx, val in enumerate(data):
            solution_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "2.3 转型路线图设计", level=2)
    
    roadmap_table = doc.add_table(rows=4, cols=4)
    roadmap_table.style = 'Table Grid'
    rm_headers = ["阶段", "时间", "核心任务", "关键里程碑"]
    for i, h in enumerate(rm_headers):
        cell = roadmap_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    roadmap = [
        ("POC验证期", "1-2月", "选择1条产线试点APS+MES", "试点产线交付+10%"),
        ("快速推广期", "3-6月", "推广至全厂、ERP集成", "全厂交付+15%"),
        ("深度优化期", "7-12月", "数据驱动决策、AI预测", "交付率达90%+"),
        ("持续运营期", "次年+", "内部能力固化、持续迭代", "交付率95%+"),
    ]
    for row_idx, data in enumerate(roadmap, 1):
        for col_idx, val in enumerate(data):
            roadmap_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "2.4 小组练习", level=2)
    doc.add_paragraph("针对鑫智制造「紧急插单处理」场景，设计完整的改善方案和数字化工具需求")

    add_heading(doc, "2.5 核心产出", level=2)
    doc.add_paragraph("[OK] 鑫智制造订单交付转型方案v1.0")
    doc.add_paragraph("[OK] 12个月数字化转型路线图")
    doc.add_paragraph("[OK] 转型价值承诺书（可作为合同附件）")

    doc.add_page_break()

    # ===== 模块三详细内容 =====
    add_heading(doc, "模块三：系统建设", level=1, color=(0, 102, 153))

    add_heading(doc, "3.1 培训目标", level=2)
    doc.add_paragraph("理解制造企业数字化系统的核心架构")
    doc.add_paragraph("掌握MES/APS/WMS等系统的选型方法")
    doc.add_paragraph("能够设计数据看板和预警规则")

    add_heading(doc, "3.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】智能制造系统架构", style='Intense Quote')
    doc.add_paragraph("")
    doc.add_paragraph("           ┌─────────────────────────────────────┐")
    doc.add_paragraph("           │           经营管理层（ERP）          │")
    doc.add_paragraph("           │     订单管理 / 物料计划 / 成本核算    │")
    doc.add_paragraph("           └─────────────┬───────────────────────┘")
    doc.add_paragraph("                         │")
    doc.add_paragraph("           ┌─────────────┴───────────────────────┐")
    doc.add_paragraph("           │           计划执行层（APS/MES）      │")
    doc.add_paragraph("           │    智能排产 / 工序跟踪 / 质量管控    │")
    doc.add_paragraph("           └─────────────┬───────────────────────┘")
    doc.add_paragraph("                         │")
    doc.add_paragraph("           ┌─────────────┴───────────────────────┐")
    doc.add_paragraph("           │           设备层（IoT/SCADA）        │")
    doc.add_paragraph("           │    设备联网 / 数据采集 / 自动控制    │")
    doc.add_paragraph("           └─────────────────────────────────────┘")

    doc.add_paragraph("【实战演练】鑫智制造数字化系统设计", style='Intense Quote')

    sys_table = doc.add_table(rows=5, cols=3)
    sys_table.style = 'Table Grid'
    sys_headers = ["系统", "核心功能", "解决痛点"]
    for i, h in enumerate(sys_headers):
        cell = sys_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    systems = [
        ("APS智能排产", "自动计算最优排产方案，考虑产能、物料、交期", "排产靠经验、紧急插单多"),
        ("MES制造执行", "工序级进度追踪、报工、质量检验", "进度不透明、信息滞后"),
        ("WMS仓库管理", "配料提醒、库存预警、先进先出", "配料不及时、库存积压"),
        ("PMC看板", "订单进度可视化、异常自动预警", "靠电话追问、问题发现晚"),
    ]
    for row_idx, data in enumerate(systems, 1):
        for col_idx, val in enumerate(data):
            sys_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "3.3 AI工具实操", level=2)
    doc.add_paragraph("使用AI工具分析历史订单数据，预测交付风险")
    doc.add_paragraph("使用AI工具自动识别生产异常模式")
    doc.add_paragraph("使用AI工具生成智能客服问答库（内部使用）")

    add_heading(doc, "3.4 核心产出", level=2)
    doc.add_paragraph("[OK] 鑫智制造数字化系统需求文档（功能清单）")
    doc.add_paragraph("[OK] 系统选型评估表（含主流产品对比）")
    doc.add_paragraph("[OK] PMC管理驾驶舱原型（Excel可操作版）")

    doc.add_page_break()

    # ===== 模块四详细内容 =====
    add_heading(doc, "模块四：变革落地", level=1, color=(0, 102, 153))

    add_heading(doc, "4.1 培训目标", level=2)
    doc.add_paragraph("掌握组织变革推动的关键方法")
    doc.add_paragraph("能够建立内部数字化运营团队")
    doc.add_paragraph("设计有效的激励机制和考核体系")

    add_heading(doc, "4.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】制造业数字化转型变革管理四步法", style='Intense Quote')
    doc.add_paragraph("第一步：领导力驱动 - 高层挂帅，成立数字化委员会")
    doc.add_paragraph("第二步：能力建设 - 培养内部数字化专员团队")
    doc.add_paragraph("第三步：流程嵌入 - 将系统使用嵌入业务流程")
    doc.add_paragraph("第四步：文化固化 - 将数据驱动写入公司价值观")

    doc.add_paragraph("【实战演练】鑫智制造变革落地方案", style='Intense Quote')

    change_table = doc.add_table(rows=5, cols=3)
    change_table.style = 'Table Grid'
    ch_headers = ["变革领域", "具体措施", "责任人"]
    for i, h in enumerate(ch_headers):
        cell = change_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    changes = [
        ("生产部门", "设立「数字化班长」岗位，KPI增加交付准时率30%权重", "生产总监"),
        ("PMC计划部", "APS使用率纳入考核，系统自动记录操作日志", "PMC经理"),
        ("IT部门", "建立系统运维SLA，7x24h响应机制", "IT总监"),
        ("人力资源", "数字化能力纳入晋升条件，开展技能认证", "HR总监"),
    ]
    for row_idx, data in enumerate(changes, 1):
        for col_idx, val in enumerate(data):
            change_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "4.3 内部讲师认证", level=2)
    doc.add_paragraph("选拔鑫智制造6名骨干成为内部数字化讲师")
    doc.add_paragraph("颁发「数字化内部认证讲师」证书")
    doc.add_paragraph("建立鑫智制造数字化知识库，持续沉淀经验")

    add_heading(doc, "4.4 核心产出", level=2)
    doc.add_paragraph("[OK] 鑫智制造数字化变革落地方案")
    doc.add_paragraph("[OK] 内部数字化运营手册")
    doc.add_paragraph("[OK] 数字化考核激励方案")

    doc.add_page_break()

    # ===== 模块五详细内容（彩蛋） =====
    add_heading(doc, "模块五（彩蛋）：价值验证", level=1, color=(0, 102, 153))

    add_heading(doc, "5.1 培训目标", level=2)
    doc.add_paragraph("理解数字化转型价值验证的方法论")
    doc.add_paragraph("能够设计并追踪转型ROI")
    doc.add_paragraph("掌握向管理层汇报转型成果的技巧")

    add_heading(doc, "5.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】转型价值验证模型", style='Intense Quote')
    doc.add_paragraph("")
    doc.add_paragraph("转型ROI = (效率提升收益 + 成本降低收益 + 质量提升收益 - 转型投入) / 转型投入 × 100%")

    doc.add_paragraph("鑫智制造价值测算示例：")
    value_table = doc.add_table(rows=6, cols=3)
    value_table.style = 'Table Grid'
    val_headers = ["改善项目", "预期收益/年", "说明"]
    for i, h in enumerate(val_headers):
        cell = value_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    values = [
        ("交付准时率 72%→95%", "+600万", "减少紧急空运、加班费、客户流失赔偿"),
        ("良品率 82%→95%", "+200万", "减少返工、报废、补料成本"),
        ("计划效率提升", "+100万", "减少计划员人工、换模时间降低"),
        ("库存周转提升", "+80万", "减少原材料和在制品库存资金占用"),
        ("合计", "+980万", "未含品牌溢价和新订单带来的增长"),
    ]
    for row_idx, data in enumerate(values, 1):
        for col_idx, val in enumerate(data):
            value_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("总投入预估：280万（系统+培训+咨询+实施）").bold = True
    doc.add_paragraph("预期ROI：250%")
    doc.add_paragraph("投资回收期：约4个月")

    add_heading(doc, "5.3 价值可视化看板", level=2)
    doc.add_paragraph("PMC管理驾驶舱：实时展示交付各项指标")
    doc.add_paragraph("价值追踪表：月度展示转型ROI趋势")
    doc.add_paragraph("管理层汇报PPT：3页精华版，直击转型价值")

    add_heading(doc, "5.4 核心产出", level=2)
    doc.add_paragraph("[OK] 鑫智制造PMC管理驾驶舱（Excel可操作版）")
    doc.add_paragraph("[OK] 月度转型价值验证报告模板")
    doc.add_paragraph("[OK] 管理层汇报PPT模板（3页精华版）")

    doc.add_page_break()

    # ===== 培训安排 =====
    add_heading(doc, "培训安排建议", level=1, color=(0, 51, 102))

    schedule_table = doc.add_table(rows=6, cols=4)
    schedule_table.style = 'Table Grid'
    sch_headers = ["模块", "建议时长", "培训形式", "产出数量"]
    for i, h in enumerate(sch_headers):
        cell = schedule_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    schedules = [
        ("模块一", "1天", "理论+实操", "3项产出"),
        ("模块二", "1天", "工作坊+演练", "3项产出"),
        ("模块三", "1天", "演示+实操", "3项产出"),
        ("模块四", "0.5天", "工作坊+讨论", "3项产出"),
        ("彩蛋模块", "0.5天", "案例+模板", "3项产出"),
    ]
    for row_idx, data in enumerate(schedules, 1):
        for col_idx, val in enumerate(data):
            schedule_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph("建议总时长：4天（可分2周执行，每周2天）")
    doc.add_paragraph("学员人数：建议20-40人/期")
    doc.add_paragraph("课后辅导：提供30天线上答疑服务")
    doc.add_paragraph("")
    doc.add_paragraph("增值服务（可选）：")
    doc.add_paragraph("  - 陪跑服务：3个月驻场辅导（额外收费）")
    doc.add_paragraph("  - 系统选型：协助对接3-5家供应商）

    doc.add_page_break()

    # ===== 附录：核心工具清单 =====
    add_heading(doc, "附录：培训核心工具清单", level=1, color=(0, 51, 102))

    tools_table = doc.add_table(rows=11, cols=3)
    tools_table.style = 'Table Grid'
    tool_headers = ["类别", "工具名称", "用途"]
    for i, h in enumerate(tool_headers):
        cell = tools_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tools = [
        ("流程设计", "Visio/Miro", "绘制业务流程图"),
        ("数据分析", "Excel/Power BI", "效率数据分析与可视化"),
        ("排产工具", "APS演示版", "智能排产演练"),
        ("文本分析", "ChatGPT/文心一言", "会议纪要、报告撰写"),
        ("项目管理", "Teambition/飞书", "任务跟踪与协作"),
        ("即时通讯", "企业微信/钉钉", "预警推送与沟通"),
        ("文档协作", "腾讯文档/飞书文档", "知识沉淀与共享"),
        ("演示汇报", "PPT模板包", "管理层汇报材料"),
        ("价值追踪", "Excel看板模板", "ROI追踪与验证"),
        ("案例沉淀", "案例库模板", "优秀案例收集整理"),
    ]
    for row_idx, data in enumerate(tools, 1):
        for col_idx, val in enumerate(data):
            tools_table.rows[row_idx].cells[col_idx].text = val

    # 保存文档
    output_path = r"d:\workboddy 培训设计\AI时代企业数字化转型实战培训体系_制造企业场景版.docx"
    doc.save(output_path)
    print(f"[OK] 文档已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    create_outline_doc()
