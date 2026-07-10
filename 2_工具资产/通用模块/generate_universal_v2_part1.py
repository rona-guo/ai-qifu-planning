# -*- coding: utf-8 -*-
"""
生成《AI时代企业数字化转型实战培训体系》通用版 - Part 1
封面+核心理念+场景设计+全景图+模块一+模块二
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return heading

def create_part1():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ===== 封面 =====
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI时代企业数字化转型实战培训体系")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— 通用版（适配多行业）")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("「商业为本 · 场景贯穿 · 成果交付 · 价值验证」")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()
    desc2 = doc.add_paragraph()
    desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc2.add_run("以企业「运营效率提升」为核心贯穿场景\n适配制造/零售/服务/科技等多行业\n\n每个模块包含：理论 + 方法论 + 案例 + 产出物")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_page_break()

    # ===== 核心理念 =====
    add_heading(doc, "一、核心理念", level=1, color=(0, 51, 102))
    doc.add_paragraph("本培训体系以「商业为本、场景贯穿、成果交付、价值验证」为核心理念。")

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run("="*55 + "\n")
    run = formula.add_run("商业成功 = 理解商业本质 × 选对商业模式 × 用好工具（数字化+AI）\n")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    run = formula.add_run("="*55)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("【重要认知】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(200, 80, 0)
    p.add_run("  数字化是'术'，AI是'术中术'，商业本质才是'道'。工具服务于商业目标，而非相反。")

    doc.add_page_break()

    # ===== 贯穿场景 =====
    add_heading(doc, "二、贯穿场景设计", level=1, color=(0, 51, 102))

    doc.add_paragraph("本培训采用「企业运营效率提升」作为贯穿场景，设计一个通用化案例供参考：")

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ["维度", "案例背景（参考）", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("企业画像", "中型企业，年营收2-10亿，员工200-1000人", "覆盖中小企业数字化转型主战场"),
        ("核心痛点", "运营效率低、数据不透明、决策靠经验", "90%企业都面临的共性问题"),
        ("改善方向", "流程优化 + 数字工具 + 数据驱动", "可适配各行业的通用路径"),
    ]
    for row_idx, row in enumerate(data, 1):
        for col_idx, val in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph("注：导师可根据学员行业替换为对应案例（制造业/零售业/服务业/科技业均有对应案例库）")

    doc.add_page_break()

    # ===== 全景图 =====
    add_heading(doc, "三、培训体系全景图", level=1, color=(0, 51, 102))

    table2 = doc.add_table(rows=6, cols=5)
    table2.style = 'Table Grid'
    headers2 = ["模块", "主题", "理论", "方法论", "核心产出"]
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    mods = [
        ("模块一", "商业诊断", "企业诊断理论", "诊断三板斧", "诊断报告"),
        ("模块二", "方案设计", "转型规划理论", "方案设计三步法", "转型方案"),
        ("模块三", "系统建设", "系统架构理论", "系统选型三问法", "需求文档"),
        ("模块四", "变革落地", "变革管理理论", "变革四步法", "落地方案"),
        ("彩蛋", "价值验证", "ROI验证理论", "价值追踪法", "验证报告"),
    ]
    for row_idx, row in enumerate(mods, 1):
        for col_idx, val in enumerate(row):
            table2.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph("建议总时长：4天（可分2周执行，每周2天） | 学员人数：20-40人/期")

    doc.add_page_break()

    # ===== 模块一 =====
    add_heading(doc, "四、模块一：商业诊断", level=1, color=(0, 102, 153))
    doc.add_paragraph("时长：1天 | 目标：掌握企业现状诊断方法，识别转型机会")

    # 理论层
    p = doc.add_paragraph()
    run = p.add_run("【理论层】企业诊断基础理论")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 153)

    theory1 = [
        "1.1 企业诊断的定义与价值",
        "  - 定义：系统性地识别企业经营问题的过程",
        "  - 价值：避免"头痛医头脚痛医脚"，找到根本原因",
        "  - 时机：新业务启动前、业绩下滑时、数字化转型前",
        "",
        "1.2 企业效率的三层分析框架",
        "  - 战略层：商业模式是否成立？价值主张是否清晰？",
        "  - 运营层：核心流程是否高效？部门协同是否顺畅？",
        "  - 执行层：员工能力是否匹配？激励机制是否有效？",
        "",
        "1.3 数字化转型的诊断视角",
        "  - 哪些环节可以通过数字化提升效率？",
        "  - 哪些数据还没有被采集和利用？",
        "  - 哪些决策还依赖经验而非数据？",
    ]
    for line in theory1:
        doc.add_paragraph(line)

    # 方法论层
    p = doc.add_paragraph()
    run = p.add_run("【方法论层】诊断三板斧")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 153, 76)

    method1 = [
        "",
        "第一斧：商业模式画布分析（60分钟）",
        "  - 工具：商业模式画布（9宫格）",
        "  - 操作：逐项填写，画出企业商业模式的"全景图"",
        "  - 输出：识别商业模式中的数字化机会点",
        "",
        "第二斧：价值链诊断（90分钟）",
        "  - 工具：价值链分析图",
        "  - 操作：绘制企业核心价值链",
        "  - 输出：识别各环节的效率损耗点",
        "",
        "第三斧：根因分析五问法（60分钟）",
        "  - 工具：5Why分析法 + 鱼骨图",
        "  - 操作：对每个痛点追问5层"为什么"",
        "  - 输出：找到问题的根本原因",
    ]
    for line in method1:
        doc.add_paragraph(line)

    # 案例层
    p = doc.add_paragraph()
    run = p.add_run("【案例层】多行业诊断案例")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(180, 80, 0)

    cases1 = [
        "",
        "案例A：制造业——某零部件厂效率诊断",
        "  背景：年营收3亿，订单交付准时率仅68%",
        "  诊断发现：订单变更频繁（35%订单有变更）是根本原因",
        "  结论：首要机会是「订单变更管控」",
        "",
        "案例B：零售业——某连锁超市诊断",
        "  背景：50家门店，客户复购率逐年下降",
        "  诊断发现：会员数据未有效利用，不了解客户需求",
        "  结论：首要机会是「会员数字化运营」",
        "",
        "案例C：服务业——某物业公司诊断",
        "  背景：500人规模，客户投诉率居高不下",
        "  诊断发现：巡检靠纸质记录，信息不透明",
        "  结论：首要机会是「巡检数字化」",
    ]
    for line in cases1:
        doc.add_paragraph(line)

    # 产出层
    p = doc.add_paragraph()
    run = p.add_run("【产出层】本模块产出物")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 0, 128)

    outputs1 = ["", "产出物1：企业商业模式画布（模板可带走）", "产出物2：企业价值链分析图", "产出物3：诊断问题清单（TOP10）", "产出物4：根因分析报告", "产出物5：数字化转型机会地图"]
    for line in outputs1:
        doc.add_paragraph(line)

    doc.add_page_break()

    # ===== 模块二 =====
    add_heading(doc, "五、模块二：方案设计", level=1, color=(0, 102, 153))
    doc.add_paragraph("时长：1天 | 目标：设计完整的数字化转型方案和路线图")

    # 理论层
    p = doc.add_paragraph()
    run = p.add_run("【理论层】转型方案设计基础理论")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 153)

    theory2 = [
        "2.1 转型方案的构成要素",
        "  - 业务目标：明确要解决什么问题，达到什么效果",
        "  - 解决方案：流程优化 + 数字工具 + 组织配套",
        "  - 实施路径：分阶段推进，而非一步到位",
        "  - 资源需求：人员、资金、时间、技术",
        "",
        "2.2 数字化工具的选择逻辑",
        "  - 效率型工具：提升现有流程的效率",
        "  - 赋能型工具：为员工提供更好的工作手段",
        "  - 创新型工具：创造新的业务模式或收入来源",
        "",
        "2.3 转型路线图设计原则",
        "  - POC优先：先小范围验证，再大规模推广",
        "  - 速赢导向：选择快速见效的场景建立信心",
    ]
    for line in theory2:
        doc.add_paragraph(line)

    # 方法论层
    p = doc.add_paragraph()
    run = p.add_run("【方法论层】方案设计三步法")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 153, 76)

    method2 = [
        "",
        "第一步：锁定核心问题（60分钟）",
        "  - 工具：痛点优先级矩阵（影响度 × 紧迫度）",
        "  - 操作：从TOP10问题中筛选TOP3核心机会",
        "",
        "第二步：设计解决方案（120分钟）",
        "  - 格式：每个问题设计"流程优化 + 数字工具 + 组织配套"组合",
        "",
        "第三步：承诺价值目标（60分钟）",
        "  - 工具：价值承诺书模板",
        "  - 操作：量化每个改善项的预期收益，明确验收标准",
    ]
    for line in method2:
        doc.add_paragraph(line)

    # 案例层
    p = doc.add_paragraph()
    run = p.add_run("【案例层】多行业转型方案案例")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(180, 80, 0)

    cases2 = [
        "",
        "案例A：制造业——订单交付效率提升方案",
        "  核心问题：订单变更管控不力",
        "  解决方案：流程优化（ECN变更管控）+ 数字工具（PLM变更模块）+ 组织配套（变更评审委员会）",
        "  价值承诺：变更频率降低50%，交付准时率提升20%，ROI 300%+",
        "",
        "案例B：零售业——会员数字化运营方案",
        "  核心问题：会员数据未有效利用",
        "  解决方案：流程优化（会员标签体系）+ 数字工具（CRM会员分析）+ 组织配套（运营专员）",
        "  价值承诺：复购率提升15%，客单价提升10%，ROI 200%+",
        "",
        "案例C：服务业——巡检数字化方案",
        "  核心问题：纸质记录，异常发现滞后",
        "  解决方案：流程优化（标准化巡检路径）+ 数字工具（移动巡检APP）+ 组织配套（培训上岗）",
        "  价值承诺：覆盖率100%，响应时间降低80%，ROI 400%+",
    ]
    for line in cases2:
        doc.add_paragraph(line)

    # 产出层
    p = doc.add_paragraph()
    run = p.add_run("【产出层】本模块产出物")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(128, 0, 128)

    outputs2 = ["", "产出物1：TOP3核心问题分析报告", "产出物2：数字化转型方案v1.0", "产出物3：12个月转型路线图", "产出物4：转型价值承诺书"]
    for line in outputs2:
        doc.add_paragraph(line)

    doc.add_page_break()

    output_path = r"d:\workboddy 培训设计\AI时代企业数字化转型实战培训体系_通用版_Part1.docx"
    doc.save(output_path)
    print(f"[OK] Part1已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    create_part1()
