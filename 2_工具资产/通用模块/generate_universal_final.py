# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return heading

def add_layer(doc, name, color, lines):
    p = doc.add_paragraph()
    run = p.add_run(f"【{name}】")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*color)
    for line in lines:
        doc.add_paragraph(line)

def create_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ===== 封面 =====
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI时代企业数字化转型实战培训体系")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("—— 通用版（适配多行业）")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("「商业为本 · 场景贯穿 · 成果交付 · 价值验证」\n以企业[运营效率提升]为核心场景\n适配制造/零售/服务/科技等多行业\n\n每个模块包含：理论 + 方法论 + 案例 + 产出物")
    run.font.size = Pt(12)

    doc.add_page_break()

    # ===== 核心理念 =====
    add_heading(doc, "一、核心理念", level=1, color=(0, 51, 102))
    doc.add_paragraph("本培训体系以「商业为本、场景贯穿、成果交付、价值验证」为核心理念。")

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run("="*50 + "\n")
    run = formula.add_run("商业成功 = 理解商业本质 x 选对商业模式 x 用好工具（数字化+AI）\n")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    run = formula.add_run("="*50)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("【重要认知】")
    run.font.bold = True
    run.font.color.rgb = RGBColor(200, 80, 0)
    p.add_run("  数字化是[术]，AI是[术中术]，商业本质才是[道]。工具服务于商业目标，而非相反。")

    doc.add_page_break()

    # ===== 贯穿场景 =====
    add_heading(doc, "二、贯穿场景设计", level=1, color=(0, 51, 102))
    doc.add_paragraph("本培训采用[企业运营效率提升]作为贯穿场景，设计一个通用化案例供参考：")

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ["维度", "案例背景（参考）", "说明"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("企业画像", "中型企业，年营收2-10亿，员工200-1000人", "覆盖中小企业数字化转型主战场"),
        ("核心痛点", "运营效率低、数据不透明、决策靠经验", "90%企业都面临的共性问题"),
        ("改善方向", "流程优化 + 数字工具 + 数据驱动", "可适配各行业的通用路径"),
    ]
    for row_idx, row in enumerate(data, 1):
        for col_idx, val in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph("注：导师可根据学员行业替换为对应案例（制造业/零售业/服务业/科技业均有对应案例库）")

    doc.add_page_break()

    # ===== 全景图 =====
    add_heading(doc, "三、培训体系全景图", level=1, color=(0, 51, 102))

    table2 = doc.add_table(rows=6, cols=5)
    table2.style = 'Table Grid'
    headers2 = ["模块", "主题", "理论", "方法论", "核心产出"]
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    mods = [
        ("模块一", "商业诊断", "企业诊断理论", "诊断三板斧", "诊断报告"),
        ("模块二", "方案设计", "转型规划理论", "方案设计三步法", "转型方案"),
        ("模块三", "系统建设", "系统架构理论", "系统选型三问法", "需求文档"),
        ("模块四", "变革落地", "变革管理理论", "变革四步法", "落地方案"),
        ("彩蛋", "价值验证", "ROI验证理论", "价值追踪法", "验证报告"),
    ]
    for row_idx, row in enumerate(mods, 1):
        for col_idx, val in enumerate(row):
            table2.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph("建议总时长：4天（可分2周执行，每周2天） | 学员人数：20-40人/期")

    doc.add_page_break()

    # ===== 模块一 =====
    add_heading(doc, "四、模块一：商业诊断", level=1, color=(0, 102, 153))
    doc.add_paragraph("时长：1天 | 目标：掌握企业现状诊断方法，识别转型机会")

    add_layer(doc, "理论层", (0, 102, 153), [
        "",
        "1.1 企业诊断的定义与价值",
        "  - 定义：系统性地识别企业经营问题的过程",
        "  - 价值：避免[头痛医头脚痛医脚]，找到根本原因",
        "  - 时机：新业务启动前、业绩下滑时、数字化转型前",
        "",
        "1.2 企业效率的三层分析框架",
        "  - 战略层：商业模式是否成立？价值主张是否清晰？",
        "  - 运营层：核心流程是否高效？部门协同是否顺畅？",
        "  - 执行层：员工能力是否匹配？激励机制是否有效？",
        "",
        "1.3 数字化转型的诊断视角",
        "  - 哪些环节可以通过数字化提升效率？",
        "  - 哪些数据还没有被采集和利用？",
        "  - 哪些决策还依赖经验而非数据？",
    ])

    add_layer(doc, "方法论层", (0, 153, 76), [
        "",
        "第一斧：商业模式画布分析（60分钟）",
        "  - 工具：商业模式画布（9宫格）",
        "  - 操作：逐项填写，画出企业商业模式的[全景图]",
        "  - 输出：识别商业模式中的数字化机会点",
        "",
        "第二斧：价值链诊断（90分钟）",
        "  - 工具：价值链分析图",
        "  - 操作：绘制企业核心价值链",
        "  - 输出：识别各环节的效率损耗点",
        "",
        "第三斧：根因分析五问法（60分钟）",
        "  - 工具：5Why分析法 + 鱼骨图",
        "  - 操作：对每个痛点追问5层[为什么]",
        "  - 输出：找到问题的根本原因",
    ])

    add_layer(doc, "案例层", (180, 80, 0), [
        "",
        "案例A：制造业——某零部件厂效率诊断",
        "  背景：年营收3亿，订单交付准时率仅68%",
        "  诊断发现：订单变更频繁（35%订单有变更）是根本原因",
        "  结论：首要机会是[订单变更管控]",
        "",
        "案例B：零售业——某连锁超市诊断",
        "  背景：50家门店，客户复购率逐年下降",
        "  诊断发现：会员数据未有效利用，不了解客户需求",
        "  结论：首要机会是[会员数字化运营]",
        "",
        "案例C：服务业——某物业公司诊断",
        "  背景：500人规模，客户投诉率居高不下",
        "  诊断发现：巡检靠纸质记录，信息不透明",
        "  结论：首要机会是[巡检数字化]",
    ])

    add_layer(doc, "产出层", (128, 0, 128), [
        "",
        "产出物1：企业商业模式画布（模板可带走）",
        "产出物2：企业价值链分析图",
        "产出物3：诊断问题清单（TOP10）",
        "产出物4：根因分析报告",
        "产出物5：数字化转型机会地图",
    ])

    doc.add_page_break()

    # ===== 模块二 =====
    add_heading(doc, "五、模块二：方案设计", level=1, color=(0, 102, 153))
    doc.add_paragraph("时长：1天 | 目标：设计完整的数字化转型方案和路线图")

    add_layer(doc, "理论层", (0, 102, 153), [
        "",
        "2.1 转型方案的构成要素",
        "  - 业务目标：明确要解决什么问题，达到什么效果",
        "  - 解决方案：流程优化 + 数字工具 + 组织配套",
        "  - 实施路径：分阶段推进，而非一步到位",
        "  - 资源需求：人员、资金、时间、技术",
        "",
        "2.2 数字化工具的选择逻辑",
        "  - 效率型工具：提升现有流程的效率",
        "  - 赋能型工具：为员工提供更好的工作手段",
        "  - 创新型工具：创造新的业务模式或收入来源",
        "",
        "2.3 转型路线图设计原则",
        "  - POC优先：先小范围验证，再大规模推广",
        "  - 速赢导向：选择快速见效的场景建立信心",
    ])

    add_layer(doc, "方法论层", (0, 153, 76), [
        "",
        "第一步：锁定核心问题（60分钟）",
        "  - 工具：痛点优先级矩阵（影响度 x 紧迫度）",
        "  - 操作：从TOP10问题中筛选TOP3核心机会",
        "",
        "第二步：设计解决方案（120分钟）",
        "  - 格式：每个问题设计[流程优化 + 数字工具 + 组织配套]组合",
        "",
        "第三步：承诺价值目标（60分钟）",
        "  - 工具：价值承诺书模板",
        "  - 操作：量化每个改善项的预期收益，明确验收标准",
    ])

    add_layer(doc, "案例层", (180, 80, 0), [
        "",
        "案例A：制造业——订单交付效率提升方案",
        "  核心问题：订单变更管控不力",
        "  解决方案：流程优化（ECN变更管控）+ 数字工具（PLM变更模块）+ 组织配套（变更评审委员会）",
        "  价值承诺：变更频率降低50%，交付准时率提升20%，ROI 300%+",
        "",
        "案例B：零售业——会员数字化运营方案",
        "  核心问题：会员数据未有效利用",
        "  解决方案：流程优化（会员标签体系）+ 数字工具（CRM会员分析）+ 组织配套（运营专员）",
        "  价值承诺：复购率提升15%，客单价提升10%，ROI 200%+",
        "",
        "案例C：服务业——巡检数字化方案",
        "  核心问题：纸质记录，异常发现滞后",
        "  解决方案：流程优化（标准化巡检路径）+ 数字工具（移动巡检APP）+ 组织配套（培训上岗）",
        "  价值承诺：覆盖率100%，响应时间降低80%，ROI 400%+",
    ])

    add_layer(doc, "产出层", (128, 0, 128), [
        "",
        "产出物1：TOP3核心问题分析报告",
        "产出物2：数字化转型方案v1.0",
        "产出物3：12个月转型路线图",
        "产出物4：转型价值承诺书",
    ])

    doc.add_page_break()

    # ===== 模块三 =====
    add_heading(doc, "六、模块三：系统建设", level=1, color=(0, 102, 153))
    doc.add_paragraph("时长：1天 | 目标：掌握数字化系统架构设计和选型方法")

    add_layer(doc, "理论层", (0, 102, 153), [
        "",
        "3.1 企业数字化系统的三层架构",
        "  - 业务应用层：ERP/CRM/MES/SCM等业务系统",
        "  - 数据中台层：主数据管理、数据仓库、API网关",
        "  - 基础设施层：云服务器、网络安全、数据存储",
        "",
        "3.2 主流数字化工具分类",
        "  - 协同办公类：钉钉/飞书/企业微信",
        "  - 经营管理类：用友/金蝶/SAP",
        "  - 生产执行类：西门子/鼎捷/慧程",
        "  - 数据分析类：帆软/Power BI/Tableau",
        "  - AI工具类：ChatGPT/文心一言/通义",
        "",
        "3.3 系统集成的关键原则",
        "  - 数据互联：打破信息孤岛，实现数据互通",
        "  - 流程贯通：端到端流程自动化",
        "  - 用户体验：系统是为用户服务的，好用优先",
    ])

    add_layer(doc, "方法论层", (0, 153, 76), [
        "",
        "第一问：这个系统解决什么业务问题？",
        "  - 不要为了上系统而上系统",
        "  - 先明确要解决的问题，再选系统",
        "",
        "第二问：这个系统能被员工真正使用吗？",
        "  - 易用性 > 功能性",
        "  - 再好的系统，没人用等于零",
        "",
        "第三问：这个系统能和企业现有系统集成吗？",
        "  - 数据孤岛是数字化最大敌人",
        "  - 选型时必须评估集成能力",
    ])

    add_layer(doc, "案例层", (180, 80, 0), [
        "",
        "案例A：制造业——MES系统选型过程",
        "  需求：需要实时掌握车间生产进度",
        "  第一问：解决什么问题？-> 生产进度可视化",
        "  第二问：员工会用吗？-> 选择移动端友好的产品",
        "  第三问：能和ERP集成吗？-> 提供标准API接口",
        "  最终选型：某国产MES，支持移动端，与用友ERP无缝集成",
        "",
        "案例B：零售业——CRM系统选型过程",
        "  需求：需要统一管理全渠道客户数据",
        "  第一问：解决什么问题？-> 客户数据分散，无法分析",
        "  第二问：店员会用吗？-> 选择操作简单的移动端",
        "  第三问：能和电商系统集成吗？-> 支持多平台数据对接",
        "  最终选型：某SaaS CRM，打通天猫/京东/线下数据",
        "",
        "案例C：服务业——巡检系统选型过程",
        "  需求：需要实时掌握巡检状态，及时发现异常",
        "  第一问：解决什么问题？-> 纸质记录无法追溯",
        "  第二问：保安会用吗？-> 选择极简操作界面",
        "  第三问：能和物业系统集成吗？-> 支持工单推送",
        "  最终选型：某物业巡检小程序，一键拍照上传",
    ])

    add_layer(doc, "AI工具实操", (100, 100, 180), [
        "",
        "现场演示：",
        "  - 使用AI工具分析历史订单数据，预测交付风险",
        "  - 使用AI工具自动生成会议纪要",
        "  - 使用AI工具构建智能问答知识库",
        "  - 使用AI工具辅助撰写系统需求文档",
    ])

    add_layer(doc, "产出层", (128, 0, 128), [
        "",
        "产出物1：企业数字化系统需求文档（功能清单）",
        "产出物2：系统选型评估表（含评估维度）",
        "产出物3：管理驾驶舱原型（Excel可操作版）",
    ])

    doc.add_page_break()

    # ===== 模块四 =====
    add_heading(doc, "七、模块四：变革落地", level=1, color=(0, 102, 153))
    doc.add_paragraph("时长：0.5天 | 目标：掌握组织变革推动方法，建立内部运营体系")

    add_layer(doc, "理论层", (0, 102, 153), [
        "",
        "4.1 数字化转型失败的常见原因",
        "  - 技术上了，业务没跟上",
        "  - 系统有了，员工不用",
        "  - 领导推了，文化没变",
        "",
        "4.2 变革管理的冰山模型",
        "  - 冰山之上：流程、系统、工具",
        "  - 冰山之下：组织结构、考核激励、文化价值观",
        "  - 真正的变革必须触及冰山之下",
        "",
        "4.3 变革成功的关键要素",
        "  - 高层承诺：一把手亲自挂帅",
        "  - 业务驱动：解决真实业务问题",
        "  - 员工参与：让员工参与设计和决策",
        "  - 持续跟进：建立长效机制，而非一阵风",
    ])

    add_layer(doc, "方法论层", (0, 153, 76), [
        "",
        "第一步：领导力驱动",
        "  - 高层挂帅，成立数字化委员会",
        "  - 明确一把手工程，由CEO或COO亲自推动",
        "",
        "第二步：能力建设",
        "  - 培养内部数字化专员团队",
        "  - 建立数字化技能认证体系",
        "",
        "第三步：流程嵌入",
        "  - 将数字化工具使用嵌入业务流程",
        "  - 建立系统使用SLA和考核机制",
        "",
        "第四步：文化固化",
        "  - 将数据驱动写入公司价值观",
        "  - 建立经验分享和复盘机制",
    ])

    add_layer(doc, "案例层", (180, 80, 0), [
        "",
        "案例：某制造企业MES上线推行过程",
        "  挑战：员工抵触，觉得增加了工作量",
        "  应对：",
        "    - 领导力：CEO亲自站台，提出[不用MES就是落后产能]",
        "    - 能力建设：选拔20名[数字化先锋]，先行培训和使用",
        "    - 流程嵌入：将MES使用纳入班组长KPI，占比30%",
        "    - 文化固化：每月评选[MES使用之星]，公开表彰",
        "  结果：3个月内系统使用率达95%，员工从抵触到主动提建议",
    ])

    add_layer(doc, "内部讲师认证", (100, 100, 180), [
        "",
        "选拔学员成为企业内部数字化讲师",
        "颁发[数字化内部认证讲师]证书",
        "建立企业内部数字化知识库",
    ])

    add_layer(doc, "产出层", (128, 0, 128), [
        "",
        "产出物1：企业数字化变革落地方案",
        "产出物2：内部数字化运营手册",
        "产出物3：数字化考核激励方案",
    ])

    doc.add_page_break()

    # ===== 彩蛋模块 =====
    add_heading(doc, "八、彩蛋模块：价值验证", level=1, color=(0, 102, 153))
    doc.add_paragraph("时长：0.5天 | 目标：掌握转型价值验证方法，量化转型成果")

    add_layer(doc, "理论层", (0, 102, 153), [
        "",
        "8.1 为什么要验证转型价值？",
        "  - 证明数字化转型的投资回报",
        "  - 持续优化，识别下一阶段的改进机会",
        "  - 为管理层决策提供数据支撑",
        "",
        "8.2 转型价值的分类",
        "  - 效率提升收益：流程自动化节省的时间",
        "  - 成本降低收益：减少人力、降低错误率",
        "  - 收入增长收益：提升销售、复购率、客户满意度",
        "",
        "8.3 转型ROI计算公式",
        "  ROI = (效率提升 + 成本降低 + 收入增长 - 转型投入) / 转型投入 x 100%",
    ])

    add_layer(doc, "方法论层", (0, 153, 76), [
        "",
        "第一步：定义价值指标",
        "  - 指标类型：过程指标（系统使用率）+ 结果指标（效率提升）",
        "",
        "第二步：建立追踪机制",
        "  - 频率：每周/月/季度追踪",
        "  - 工具：管理驾驶舱实时看板",
        "",
        "第三步：验证与汇报",
        "  - 月度验证：对比实际收益与承诺收益",
        "  - 管理层汇报：用数据说话，用图表呈现",
    ])

    add_layer(doc, "案例层", (180, 80, 0), [
        "",
        "案例：某制造企业MES上线6个月价值验证",
        "  投入：系统150万 + 培训20万 + 实施30万 = 200万",
        "  收益验证：",
        "    - 生产效率提升15% -> 节省加班费50万/年",
        "    - 良品率提升5% -> 减少报废30万/年",
        "    - 在制品库存降低20% -> 释放资金80万",
        "  6个月ROI = (50+30+80)/200 = 80%",
        "  预计年ROI = 160%，投资回收期8个月",
    ])

    add_layer(doc, "产出层", (128, 0, 128), [
        "",
        "产出物1：企业价值追踪指标矩阵",
        "产出物2：月度价值验证报告模板",
        "产出物3：管理驾驶舱模板（Excel版）",
        "产出物4：管理层汇报PPT模板（3页精华版）",
    ])

    doc.add_page_break()

    # ===== 附录 =====
    add_heading(doc, "九、附录：核心工具清单", level=1, color=(0, 51, 102))

    tools_table = doc.add_table(rows=11, cols=3)
    tools_table.style = 'Table Grid'
    headers = ["类别", "工具名称", "用途"]
    for i, h in enumerate(headers):
        cell = tools_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    tools = [
        ("商业模式", "Canvanizer", "绘制商业模式画布"),
        ("流程设计", "Visio/Miro", "绘制业务流程图"),
        ("数据分析", "Excel/Power BI", "数据分析与可视化"),
        ("AI工具", "ChatGPT/文心一言", "内容生成、辅助分析"),
        ("协同办公", "飞书/钉钉/企业微信", "日常沟通、任务管理"),
        ("文档协作", "腾讯文档/飞书文档", "知识沉淀与共享"),
        ("原型设计", "Axure/Figma", "系统原型设计"),
        ("演示汇报", "PPT模板包", "管理层汇报材料"),
        ("价值追踪", "Excel看板模板", "ROI追踪与验证"),
        ("案例管理", "案例库模板", "优秀案例收集整理"),
    ]
    for row_idx, data in enumerate(tools, 1):
        for col_idx, val in enumerate(data):
            tools_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    # ===== 总结 =====
    add_heading(doc, "十、培训安排总结", level=1, color=(0, 51, 102))

    summary_table = doc.add_table(rows=7, cols=4)
    summary_table.style = 'Table Grid'
    headers = ["模块", "时长", "核心方法论", "核心产出"]
    for i, h in enumerate(headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    summary = [
        ("模块一", "1天", "诊断三板斧", "诊断报告"),
        ("模块二", "1天", "方案设计三步法", "转型方案"),
        ("模块三", "1天", "系统选型三问法", "需求文档"),
        ("模块四", "0.5天", "变革四步法", "落地方案"),
        ("彩蛋", "0.5天", "价值追踪三步法", "验证报告"),
        ("合计", "4天", "-", "15项产出物"),
    ]
    for row_idx, data in enumerate(summary, 1):
        for col_idx, val in enumerate(data):
            cell = summary_table.rows[row_idx].cells[col_idx]
            cell.text = val
            if row_idx == 6:
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    doc.add_paragraph("【服务说明】")
    doc.add_paragraph("  - 标准版：4天（完整体系）")
    doc.add_paragraph("  - 紧凑版：2天（聚焦核心模块）")
    doc.add_paragraph("  - 课后辅导：30天线上答疑服务")
    doc.add_paragraph("")
    doc.add_paragraph("【增值服务】（可选，额外收费）")
    doc.add_paragraph("  - 陪跑服务：3-6个月驻场辅导")
    doc.add_paragraph("  - 系统选型：协助对接供应商")
    doc.add_paragraph("  - 行业定制：根据行业特点补充案例")

    output_path = r"d:\workboddy 培训设计\AI时代企业数字化转型实战培训体系_通用版_丰富版.docx"
    doc.save(output_path)
    print(f"[OK] 文档已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    create_doc()
