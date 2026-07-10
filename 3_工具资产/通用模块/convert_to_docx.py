"""
Convert the V3.0 whitepaper MD to a formatted DOCX file.
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_from_md(doc, header_row, data_rows):
    """Add a table from parsed MD table data."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = cell_text.strip()
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_background(cell, '4472C4')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    # Data rows
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx + 1].cells[col_idx]
                # Handle bold markers
                text = cell_text.strip()
                cell.text = text.replace('**', '')
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
    
    doc.add_paragraph()

def process_inline_formatting(paragraph, text):
    """Process bold and other inline formatting."""
    # Split by ** for bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle `code` formatting
            code_parts = re.split(r'(`.*?`)', part)
            for cp in code_parts:
                if cp.startswith('`') and cp.endswith('`'):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    if cp:
                        paragraph.add_run(cp)

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                # Add light gray background
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'F5F5F5')
                shading.set(qn('w:val'), 'clear')
                p._p.get_or_add_pPr().append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Table detection
        if '|' in line and i + 1 < len(lines) and re.match(r'^[\s\|:-]+$', lines[i+1].strip()):
            # Parse table
            header_cells = [c.strip() for c in line.split('|')[1:-1]]
            i += 2  # Skip header and separator
            data_rows = []
            while i < len(lines) and '|' in lines[i].strip() and lines[i].strip():
                row_cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                data_rows.append(row_cells)
                i += 1
            add_table_from_md(doc, header_cells, data_rows)
            continue
        
        # Headings
        if line.startswith('# '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line.startswith('## '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line.startswith('### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=3)
            heading.style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif line.startswith('#### '):
            text = line[5:].strip()
            heading = doc.add_heading(text, level=4)
            heading.style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        # Horizontal rule
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('─' * 40)
            run.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)
        # Blockquote
        elif line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            process_inline_formatting(p, text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                run.font.size = Pt(10)
            # Left border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '4472C4')
            pBdr.append(left)
            pPr.append(pBdr)
        # List items
        elif re.match(r'^[\s]*[-*]\s', line):
            text = re.sub(r'^[\s]*[-*]\s', '', line)
            level = (len(line) - len(line.lstrip())) // 2
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
            process_inline_formatting(p, text.strip())
            for run in p.runs:
                run.font.size = Pt(10)
        elif re.match(r'^[\s]*\d+\.\s', line):
            text = re.sub(r'^[\s]*\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            process_inline_formatting(p, text.strip())
        # Empty line
        elif not line.strip():
            pass
        # Normal paragraph
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            process_inline_formatting(p, line)
            for run in p.runs:
                run.font.size = Pt(10.5)
        
        i += 1
    
    doc.save(docx_path)
    print(f'DOCX saved: {docx_path}')

# Execute
md_path = r'D:\AI落地\2_产品设计\服务产品\标准企业AI落地解决方案_V3.0白皮书版.md'
docx_path = r'D:\AI落地\2_产品设计\服务产品\标准企业AI落地解决方案_V3.0白皮书版.docx'
convert_md_to_docx(md_path, docx_path)
