# -*- coding: utf-8 -*-
"""
生成《AI时代企业数字化转型实战培训体系》CEM客户体验管理场景版
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1, color=None):
    """添加标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return heading

def create_outline_doc():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)

    # ===== 封面页 =====
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI时代企业数字化转型实战培训体系")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("—— CEM客户体验管理场景版")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()
    doc.add_paragraph()

    # 副标题说明
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("以华东某智能家居制造商为贯穿案例\n从客户体验诊断到价值验证的完整培训体系")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ===== 核心理念 =====
    add_heading(doc, "核心理念", level=1, color=(0, 51, 102))
    
    doc.add_paragraph("本培训体系以「商业为本、场景贯穿、成果交付、价值验证」为核心理念，采用CEM客户体验管理作为贯穿场景，让学员在真实企业案例中掌握数字化转型的完整方法论。")

    # 核心公式框
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = formula.add_run("\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n")
    run.font.size = Pt(12)
    run = formula.add_run("商业成功 = 理解商业本质 × 选对商业模式 × 用好工具（数字化+AI）\n")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    run = formula.add_run("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n")
    run.font.size = Pt(12)

    doc.add_paragraph()

    # ===== 案例企业背景 =====
    add_heading(doc, "贯穿案例：华东某智能家居制造商", level=1, color=(0, 51, 102))

    # 企业信息表格
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    info_data = [
        ("企业名称", "华东某智能家居制造商（化名：智居科技）"),
        ("企业规模", "年营收8亿元，员工1200人，研发人员200人"),
        ("主营产品", "智能门锁、智能开关、全屋智能控制系统"),
        ("渠道结构", "电商40%、经销商35%、工程渠道25%"),
        ("当前痛点", "客户投诉率18%、NPS仅32分、复购率仅22%"),
        ("培训目标", "NPS提升至55分、复购率提升至40%、投诉率降至5%"),
    ]

    for i, (label, value) in enumerate(info_data):
        cell = table.rows[i].cells[0]
        cell.text = label
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "E8F4FD")
        table.rows[i].cells[1].text = value

    doc.add_paragraph()

    # ===== 培训体系全景图 =====
    add_heading(doc, "培训体系全景图：五环驱动模型", level=1, color=(0, 51, 102))

    # 五环模型说明
    doc.add_paragraph("本培训体系采用「五环驱动」模型，以CEM客户体验管理为核心，贯穿从诊断到验证的完整闭环：")

    modules = [
        ("模块一", "客户体验诊断", "客户旅程地图绘制", "发现体验断点"),
        ("模块二", "CEM策略设计", "体验指标体系构建", "找准优化方向"),
        ("模块三", "体验系统建设", "数字化平台搭建", "实现数据闭环"),
        ("模块四", "变革落地执行", "组织能力升级", "推动持续改进"),
        ("模块五（彩蛋）", "价值验证交付", "成果可视化呈现", "验证培训价值"),
    ]

    table2 = doc.add_table(rows=6, cols=4)
    table2.style = 'Table Grid'

    headers = ["模块", "主题", "核心动作", "产出"]
    for i, h in enumerate(headers):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "003366")

    for row_idx, (m, t, a, p) in enumerate(modules, 1):
        table2.rows[row_idx].cells[0].text = m
        table2.rows[row_idx].cells[1].text = t
        table2.rows[row_idx].cells[2].text = a
        table2.rows[row_idx].cells[3].text = p

    doc.add_page_break()

    # ===== 模块一详细内容 =====
    add_heading(doc, "模块一：客户体验诊断", level=1, color=(0, 102, 153))

    add_heading(doc, "1.1 培训目标", level=2)
    doc.add_paragraph("• 掌握客户体验诊断的核心方法论")
    doc.add_paragraph("• 能够独立绘制客户旅程地图")
    doc.add_paragraph("• 识别智居科技的体验断点与机会点")

    add_heading(doc, "1.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】", style='Intense Quote')
    doc.add_paragraph("客户体验管理（CEM）基础理论")
    doc.add_paragraph("• CEM定义：系统化管理客户全生命周期体验的策略与方法")
    doc.add_paragraph("• CEM与CRM的区别：CRM管理客户关系，CEM管理客户感受")
    doc.add_paragraph("• CEM核心指标：NPS（净推荐值）、CES（客户费力指数）、CSAT（满意度）")

    doc.add_paragraph("【实战演练】智居科技客户旅程地图绘制", style='Intense Quote')

    # 客户旅程阶段表格
    journey_table = doc.add_table(rows=7, cols=4)
    journey_table.style = 'Table Grid'
    journey_headers = ["阶段", "触点", "客户行为", "体验痛点"]
    for i, h in enumerate(journey_headers):
        cell = journey_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")

    journey_data = [
        ("认知", "电商平台/社交媒体", "浏览产品详情", "信息繁杂，难以抉择"),
        ("考虑", "客服咨询/评测视频", "对比多家品牌", "客服响应慢、专业度不足"),
        ("购买", "官网/电商下单", "完成交易", "支付环节卡顿"),
        ("开箱", "收到产品", "安装体验", "安装说明复杂、配件遗漏"),
        ("使用", "日常使用", "产品体验", "APP连接不稳定、功能繁琐"),
        ("复购/推荐", "评价/分享", "口碑传播", "缺乏复购激励、推荐机制"),
    ]
    for row_idx, (s, t, b, p) in enumerate(journey_data, 1):
        journey_table.rows[row_idx].cells[0].text = s
        journey_table.rows[row_idx].cells[1].text = t
        journey_table.rows[row_idx].cells[2].text = b
        journey_table.rows[row_idx].cells[3].text = p

    doc.add_paragraph()

    add_heading(doc, "1.3 小组练习", level=2)
    doc.add_paragraph("分组绘制智居科技「智能门锁」产品的客户旅程地图，识别TOP3体验痛点")

    add_heading(doc, "1.4 核心产出", level=2)
    doc.add_paragraph("✅ 智居科技客户旅程地图（v1.0）")
    doc.add_paragraph("✅ 客户体验痛点分析报告")
    doc.add_paragraph("✅ 体验优化机会矩阵")

    doc.add_page_break()

    # ===== 模块二详细内容 =====
    add_heading(doc, "模块二：CEM策略设计", level=1, color=(0, 102, 153))

    add_heading(doc, "2.1 培训目标", level=2)
    doc.add_paragraph("• 建立完整的客户体验指标体系")
    doc.add_paragraph("• 设计可落地的触点优化方案")
    doc.add_paragraph("• 制定清晰的体验改善路线图")

    add_heading(doc, "2.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】体验指标体系设计", style='Intense Quote')
    doc.add_paragraph("• 三层指标体系：战略层（北极星指标）→ 运营层（过程指标）→ 触点层（行动指标）")
    doc.add_paragraph("• 智居科技指标体系设计：")
    doc.add_paragraph("  - 北极星指标：NPS从32分提升至55分")
    doc.add_paragraph("  - 过程指标：首次响应时长<2min、问题一次解决率>85%")
    doc.add_paragraph("  - 行动指标：各触点CSAT评分>4.5分")

    doc.add_paragraph("【实战演练】智居科技CEM改善方案设计", style='Intense Quote')

    solution_table = doc.add_table(rows=5, cols=4)
    solution_table.style = 'Table Grid'
    sol_headers = ["痛点", "根因分析", "改善策略", "预期效果"]
    for i, h in enumerate(sol_headers):
        cell = solution_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")

    solutions = [
        ("APP连接不稳定", "云平台并发能力不足", "升级云架构，引入边缘计算", "连接成功率从78%→95%"),
        ("客服响应慢", "人工接待能力饱和", "引入AI客服+智能路由", "响应时长从8min→1.5min"),
        ("复购率低", "缺乏生命周期运营", "搭建会员体系，精准营销", "复购率从22%→40%"),
        ("口碑传播弱", "缺乏推荐激励机制", "设计老带新奖励方案", "推荐转化率提升30%"),
    ]
    for row_idx, sol in enumerate(solutions, 1):
        for col_idx, val in enumerate(sol):
            solution_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "2.3 小组练习", level=2)
    doc.add_paragraph("针对智居科技「首次安装体验」触点，设计完整的改善方案")

    add_heading(doc, "2.4 核心产出", level=2)
    doc.add_paragraph("✅ 智居科技体验指标体系文档")
    doc.add_paragraph("✅ TOP10体验改善方案")
    doc.add_paragraph("✅ CEM实施路线图（12个月版本）")

    doc.add_page_break()

    # ===== 模块三详细内容 =====
    add_heading(doc, "模块三：体验系统建设", level=1, color=(0, 102, 153))

    add_heading(doc, "3.1 培训目标", level=2)
    doc.add_paragraph("• 理解CEM数字化平台的核心架构")
    doc.add_paragraph("• 掌握体验数据采集与分析方法")
    doc.add_paragraph("• 能够搭建基础的体验管理看板")

    add_heading(doc, "3.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】CEM平台架构设计", style='Intense Quote')
    doc.add_paragraph("• 数据采集层：多触点反馈收集（APP/小程序/电话/短信/邮件）")
    doc.add_paragraph("• 数据分析层：NLP语义分析 + 情感识别 + 趋势预测")
    doc.add_paragraph("• 行动触发层：工单流转 + 预警机制 + 个性化推荐")
    doc.add_paragraph("• 效果追踪层：实时看板 + 周期性报告 + 闭环验证")

    doc.add_paragraph("【实战演练】智居科技体验平台搭建", style='Intense Quote')

    platform_table = doc.add_table(rows=5, cols=3)
    platform_table.style = 'Table Grid'
    plat_headers = ["平台模块", "核心功能", "技术方案"]
    for i, h in enumerate(plat_headers):
        cell = platform_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")

    platforms = [
        ("反馈采集系统", "全渠道NPS/CSAT采集", "微信小程序+短信+邮件多触点"),
        ("VOC分析平台", "客户之声智能分析", "AI NLP语义分析+情感打分"),
        ("体验看板", "实时体验指标监控", "数据可视化+自定义配置"),
        ("预警工单系统", "体验问题自动触发", "规则引擎+企微/钉钉推送"),
    ]
    for row_idx, p in enumerate(platforms, 1):
        for col_idx, val in enumerate(p):
            platform_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    # AI工具应用
    add_heading(doc, "3.3 AI工具实操", level=2)
    doc.add_paragraph("• 使用AI工具进行客户评论情感分析（现场演示）")
    doc.add_paragraph("• 使用AI工具自动生成投诉分类标签")
    doc.add_paragraph("• 使用AI工具预测客户流失风险")

    add_heading(doc, "3.4 核心产出", level=2)
    doc.add_paragraph("✅ 智居科技CEM平台功能需求文档")
    doc.add_paragraph("✅ 体验数据采集方案")
    doc.add_paragraph("✅ 体验管理看板原型（可操作的Excel版本）")

    doc.add_page_break()

    # ===== 模块四详细内容 =====
    add_heading(doc, "模块四：变革落地执行", level=1, color=(0, 102, 153))

    add_heading(doc, "4.1 培训目标", level=2)
    doc.add_paragraph("• 掌握组织变革推动的关键方法")
    doc.add_paragraph("• 建立内部CEM运营团队")
    doc.add_paragraph("• 设计有效的激励机制")

    add_heading(doc, "4.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】CEM组织变革四步法", style='Intense Quote')
    doc.add_paragraph("• 第一步：意识唤醒——让全员理解「体验是产品」")
    doc.add_paragraph("• 第二步：能力建设——培养内部CEM专员团队")
    doc.add_paragraph("• 第三步：流程嵌入——将体验指标嵌入业务流程")
    doc.add_paragraph("• 第四步：文化固化——将「客户优先」写入公司价值观")

    doc.add_paragraph("【实战演练】智居科技变革落地方案", style='Intense Quote')

    change_table = doc.add_table(rows=5, cols=3)
    change_table.style = 'Table Grid'
    ch_headers = ["变革领域", "具体措施", "责任人"]
    for i, h in enumerate(ch_headers):
        cell = change_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")

    changes = [
        ("客服部门", "设立「体验专员」岗位，KPI增加NPS权重30%", "客服总监"),
        ("产品研发", "建立VOC需求优先级评审机制", "产品VP"),
        ("市场营销", "投放素材增加「体验元素」审核", "市场总监"),
        ("经销商", "经销商NPS纳入返利考核体系", "销售总监"),
    ]
    for row_idx, c in enumerate(changes, 1):
        for col_idx, val in enumerate(c):
            change_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()

    add_heading(doc, "4.3 内部讲师认证", level=2)
    doc.add_paragraph("• 选拔智居科技5名骨干成为内部CEM讲师")
    doc.add_paragraph("• 颁发「CEM内部认证讲师」证书")
    doc.add_paragraph("• 建立内部CEM知识库，持续沉淀")

    add_heading(doc, "4.4 核心产出", level=2)
    doc.add_paragraph("✅ 智居科技CEM变革落地方案")
    doc.add_paragraph("✅ 内部CEM运营手册")
    doc.add_paragraph("✅ 体验考核激励方案")

    doc.add_page_break()

    # ===== 模块五详细内容 =====
    add_heading(doc, "模块五（彩蛋）：价值验证交付", level=1, color=(0, 102, 153))

    add_heading(doc, "5.1 培训目标", level=2)
    doc.add_paragraph("• 理解数字化转型价值验证的方法论")
    doc.add_paragraph("• 能够设计并追踪体验改善ROI")
    doc.add_paragraph("• 掌握价值可视化呈现技巧")

    add_heading(doc, "5.2 核心内容", level=2)

    doc.add_paragraph("【理论框架】CEM价值验证模型", style='Intense Quote')
    doc.add_paragraph("体验投资回报率 = (体验提升带来的收益 - 体验改善成本) / 体验改善成本 × 100%")

    doc.add_paragraph("智居科技价值测算示例：")
    value_table = doc.add_table(rows=5, cols=3)
    value_table.style = 'Table Grid'
    val_headers = ["改善项目", "预期收益/年", "说明"]
    for i, h in enumerate(val_headers):
        cell = value_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "0066CC")

    values = [
        ("NPS 32→55分", "+800万", "复购率提升带来的营收增长"),
        ("投诉率 18%→5%", "+200万", "减少退款、售后、客诉处理成本"),
        ("客服效率提升", "+150万", "AI客服节省人力成本"),
        ("品牌口碑提升", "+300万", "老带新推荐增加的营收"),
    ]
    for row_idx, v in enumerate(values, 1):
        for col_idx, val in enumerate(v):
            value_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph("总投入预估：200万（系统+培训+咨询）")
    doc.add_paragraph("预期ROI：525%")
    doc.add_paragraph("投资回收期：约3个月")

    add_heading(doc, "5.3 价值可视化看板", level=2)
    doc.add_paragraph("• 体验仪表盘：实时展示NPS/CSAT/CES趋势")
    doc.add_paragraph("• 价值追踪表：月度展示体验改善ROI")
    doc.add_paragraph("• 案例故事库：收集客户好评、投诉处理案例")

    add_heading(doc, "5.4 核心产出", level=2)
    doc.add_paragraph("✅ 智居科技体验价值追踪看板（Excel可操作版）")
    doc.add_paragraph("✅ 月度CEM价值验证报告模板")
    doc.add_paragraph("✅ 管理层汇报PPT模板（3页精华版）")

    doc.add_page_break()

    # ===== 培训安排 =====
    add_heading(doc, "培训安排建议", level=1, color=(0, 51, 102))

    schedule_table = doc.add_table(rows=6, cols=4)
    schedule_table.style = 'Table Grid'
    sch_headers = ["模块", "建议时长", "培训形式", "产出数量"]
    for i, h in enumerate(sch_headers):
        cell = schedule_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")

    schedules = [
        ("模块一", "1天", "理论+实操", "3项产出"),
        ("模块二", "1天", "工作坊+演练", "3项产出"),
        ("模块三", "1天", "演示+实操", "3项产出"),
        ("模块四", "0.5天", "工作坊+讨论", "3项产出"),
        ("模块五（彩蛋）", "0.5天", "案例+模板", "3项产出"),
    ]
    for row_idx, s in enumerate(schedules, 1):
        for col_idx, val in enumerate(s):
            schedule_table.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph()
    doc.add_paragraph("💡 建议总时长：4天（可分2周执行，每周2天）")
    doc.add_paragraph("💡 学员人数：建议20-40人/期")
    doc.add_paragraph("💡 课后辅导：提供30天线上答疑服务")

    doc.add_page_break()

    # ===== 附录：核心工具清单 =====
    add_heading(doc, "附录：培训核心工具清单", level=1, color=(0, 51, 102))

    tools_table = doc.add_table(rows=11, cols=3)
    tools_table.style = 'Table Grid'
    tool_headers = ["类别", "工具名称", "用途"]
    for i, h in enumerate(tool_headers):
        cell = tools_table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, "003366")

    tools = [
        ("客户旅程", "Canvanizer/Miro", "绘制客户旅程地图"),
        ("数据分析", "Excel/Power BI", "体验数据分析与可视化"),
        ("问卷调查", "问卷星/腾讯问卷", "NPS/CSAT采集"),
        ("文本分析", "ChatGPT/文心一言", "客户评论情感分析"),
        ("项目管理", "Teambition/飞书", "任务跟踪与协作"),
        ("即时通讯", "企业微信/钉钉", "预警推送与沟通"),
        ("文档协作", "腾讯文档/飞书文档", "知识沉淀与共享"),
        ("演示汇报", "PPT模板包", "管理层汇报材料"),
        ("价值追踪", "Excel看板模板", "ROI追踪与验证"),
        ("案例沉淀", "案例库模板", "优秀案例收集整理"),
    ]
    for row_idx, t in enumerate(tools, 1):
        for col_idx, val in enumerate(t):
            tools_table.rows[row_idx].cells[col_idx].text = val

    # 保存文档
    output_path = r"d:\workboddy 培训设计\AI时代企业数字化转型实战培训体系_CEM场景版.docx"
    doc.save(output_path)
    print(f"[OK] 文档已生成：{output_path}")
    return output_path

if __name__ == "__main__":
    create_outline_doc()
